from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.docstore.document import Document
from langchain.docstore import InMemoryDocstore  # 追加
from docx import Document as DocxDocument
from langchain.text_splitter import CharacterTextSplitter
import chardet
import faiss
import numpy as np

def read_file_with_fallback(file_path):
    with open(file_path, 'rb') as f:
        raw_data = f.read()
        result = chardet.detect(raw_data)
        encoding = result['encoding']
    try:
        with open(file_path, encoding=encoding) as f:
            return f.read()
    except (UnicodeDecodeError, TypeError):
        for fallback_encoding in ['utf-8', 'shift_jis', 'iso-8859-1', 'cp932']:
            try:
                with open(file_path, encoding=fallback_encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue  # 次のエンコーディングを試す
    raise ValueError(f"{file_path}のファイルをデコードできませんでした。")
# HuggingFaceEmbeddingsのモデル名を指定
embedding = HuggingFaceEmbeddings(model_name="intfloat/multilingual-e5-large")
text_splitter = CharacterTextSplitter(chunk_size=300, chunk_overlap=32)
# ファイルパスのリスト
file_paths = [
    "/Users/SOL-Project/nago-workshop/src/document/アップルバナナ1.docx",
    "/Users/SOL-Project/nago-workshop/src/document/アップルバナナ2.docx",
    "/Users/SOL-Project/nago-workshop/src/document/マルエスファーム1.docx",
    "/Users/SOL-Project/nago-workshop/src/document/荘司幸一郎1.docx",
    "/Users/SOL-Project/nago-workshop/src/document/荘司幸一郎2.docx",
]
# ドキュメントリストの作成
docs = []
for file_path in file_paths:
    try:
        text = read_file_with_fallback(file_path)
        texts = text_splitter.split_text(text)
        for t in texts:
            docs.append(Document(page_content=t, metadata={"source": file_path}))
    except ValueError as e:
        print(e)  # エラーを適切に処理
# 各ドキュメントの埋め込みを計算
doc_embeddings = np.array(embedding.embed_documents([doc.page_content for doc in docs]), dtype=np.float32)
# FAISSのインデックスを作成
dimension = doc_embeddings.shape[1]  # 埋め込みの次元数
index = faiss.IndexFlatL2(dimension)  # L2距離を使ったインデックスを作成
# ドキュメントの埋め込みをインデックスに追加
index.add(doc_embeddings)
# docstoreをInMemoryDocstoreに変更
docstore = InMemoryDocstore({doc.metadata["source"]: doc for doc in docs})
# FAISSのインスタンスを作成
faiss_index = FAISS(
    embedding_function=embedding.embed_query,
    index=index,
    docstore=docstore,  # InMemoryDocstoreを使用
    index_to_docstore_id={i: doc.metadata["source"] for i, doc in enumerate(docs)}
)
# クエリの検索
query = "アップルバナナについて教えてください。"
results = faiss_index.similarity_search(query, k=2)
# 関連するファイルの内容全体を表示
for doc in results:
    result_doc = DocxDocument(doc.metadata["source"])
    file_content = "\n".join([para.text for para in result_doc.paragraphs])
    print(file_content)